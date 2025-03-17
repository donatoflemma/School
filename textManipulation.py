from docx import Document

# Creare un documento Word
doc = Document()

# Testo della lettera
doc.add_paragraph("Flemma Donato")
doc.add_paragraph("Dolgensee str 21   10319")
doc.add_paragraph("\nHo trovato molto interessante la vostra azienda e sarei entusiasta di potermi candidare per uno stage.")
doc.add_paragraph("\nCordiali saluti,\nMario Rossi")

# Salvare il file Word
doc.save("Lettera_Presentazione.docx")

